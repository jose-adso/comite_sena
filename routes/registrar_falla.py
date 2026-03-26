import os
import secrets
from datetime import datetime

from flask import render_template, request, redirect, url_for, flash, session
from docx import Document
from docx.shared import Inches

from models.database import db, bcrypt
from models.usuario import Usuario
from models.falla import Falla
from routes.auth import auth_bp
from routes.utils import (
    es_admin,
    _username_seguro,
    BCRYPT_ROUNDS,
    _replace_placeholders_doc,
    _enviar_pdf_siempre,
    etiqueta_rol_visible,
)

# Importar función de formateo de fecha desde notificaciones
from routes.notificaciones import formatear_fecha_larga


@auth_bp.route('/registrar-falla')
def vista_registrar_falla():
    """Vista para registrar una falta de aprendiz"""
    if 'usuario_id' not in session:
        flash('Debe iniciar sesión primero', 'error')
        return redirect(url_for('auth.index'))

    nombre_instructor = session.get('nombre', session.get('username', 'Instructor'))
    return render_template(
        'registrar_falla.html',
        username=session['username'],
        nombre_instructor=nombre_instructor,
        rol_visible=etiqueta_rol_visible(),
    )


@auth_bp.route('/registrar-falla', methods=['POST'])
def registrar_falla():
    """Procesar el registro de una falta de aprendiz"""
    if 'usuario_id' not in session:
        flash('Debe iniciar sesión primero', 'error')
        return redirect(url_for('auth.index'))

    nombre_instructor = (request.form.get('nombre_instructor') or '').strip()
    cedula_instructor = (request.form.get('cedula_instructor') or '').strip()
    nombre_ficha = (request.form.get('nombre_ficha') or '').strip()
    numero_ficha = (request.form.get('numero_ficha') or '').strip()
    nombre_aprendiz = (request.form.get('nombre_aprendiz') or '').strip()
    documento_aprendiz = (request.form.get('documento_aprendiz') or '').strip()
    correo_aprendiz = (request.form.get('correo_aprendiz') or '').strip()
    telefono_aprendiz = (request.form.get('telefono_aprendiz') or '').strip()
    descripcion_faltas = (request.form.get('descripcion_faltas') or '').strip()
    fecha_falta = request.form.get('fecha')

    # Corrección automática si los campos están invertidos
    if len(numero_ficha) > 20 and len(nombre_ficha) <= 20:
        nombre_ficha, numero_ficha = numero_ficha, nombre_ficha

    limites = [
        ('Nombre instructor', nombre_instructor, 100),
        ('Cédula instructor', cedula_instructor, 20),
        ('Nombre ficha', nombre_ficha, 100),
        ('Número ficha', numero_ficha, 20),
        ('Nombre aprendiz', nombre_aprendiz, 100),
        ('Documento aprendiz', documento_aprendiz, 20),
        ('Correo aprendiz', correo_aprendiz, 120),
        ('Teléfono aprendiz', telefono_aprendiz, 20),
    ]
    for etiqueta, valor, max_len in limites:
        if len(valor) > max_len:
            flash(f'{etiqueta} supera el límite ({max_len} caracteres)', 'error')
            return redirect(url_for('auth.vista_registrar_falla'))

    fecha_falta_date = datetime.strptime(fecha_falta, '%Y-%m-%d').date() if fecha_falta else None

    # Obtener usuario instructor
    usuario_local = Usuario.query.get(session.get('usuario_id'))
    if not usuario_local:
        username_sesion = session.get('username', 'instructor')
        correo_sesion = (session.get('correo') or '').strip().lower()
        usuario_local = Usuario.query.filter_by(email=correo_sesion).first() if correo_sesion else None
        if not usuario_local:
            usuario_local = Usuario.query.filter_by(username=_username_seguro(username_sesion)).first()
        if not usuario_local:
            temp_password = secrets.token_urlsafe(24)
            usuario_local = Usuario(
                username=_username_seguro(username_sesion),
                password_hash=bcrypt.generate_password_hash(temp_password, rounds=BCRYPT_ROUNDS).decode('utf-8'),
                rol='instructor',
                nombre=username_sesion,
                email=correo_sesion or None,
                debe_cambiar_password=False,
            )
            db.session.add(usuario_local)
            db.session.commit()
        session['usuario_id'] = usuario_local.id

    from app import app as flask_app

    # Manejar evidencias (múltiples archivos)
    evidencias = request.files.getlist('evidencia')
    archivos_guardados = []
    evidencia_ruta = ''

    if evidencias and evidencias[0].filename:
        upload_folder = os.path.join(flask_app.root_path, 'static', 'evidencias')
        os.makedirs(upload_folder, exist_ok=True)
        for evidencia in evidencias:
            if evidencia and evidencia.filename:
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                filename = f"{timestamp}_{evidencia.filename}"
                evidencia.save(os.path.join(upload_folder, filename))
                archivos_guardados.append(f"evidencias/{filename}")
        evidencia_ruta = ','.join(archivos_guardados)

    # Manejar firma
    firma = request.files.get('firma')
    firma_ruta = ''
    if firma and firma.filename:
        upload_folder = os.path.join(flask_app.root_path, 'static', 'firmas')
        os.makedirs(upload_folder, exist_ok=True)
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"firma_{timestamp}_{firma.filename}"
        firma.save(os.path.join(upload_folder, filename))
        firma_ruta = f"firmas/{filename}"

    try:
        falla = Falla(
            instructor_id=usuario_local.id,
            nombre_instructor=nombre_instructor,
            cedula_instructor=cedula_instructor,
            correo_instructor=session.get('correo', ''),
            nombre_ficha=nombre_ficha,
            numero_ficha=numero_ficha,
            nombre_aprendiz=nombre_aprendiz,
            documento_aprendiz=documento_aprendiz,
            correo_aprendiz=correo_aprendiz,
            telefono_aprendiz=telefono_aprendiz,
            descripcion_faltas=descripcion_faltas,
            fecha_falta=fecha_falta_date,
            evidencia=evidencia_ruta,
            firma=firma_ruta,
        )
        db.session.add(falla)
        db.session.commit()
        flash('Falla registrada exitosamente', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Error al registrar la falla: {str(e)}', 'error')

    return redirect(url_for('auth.vista_registrar_falla'))


@auth_bp.route('/generar-formato/<int:falla_id>')
def generar_formato(falla_id):
    """Generar formato PDF de comité para una falla"""
    if 'usuario_id' not in session:
        flash('Debe iniciar sesión primero', 'error')
        return redirect(url_for('auth.index'))

    falla = Falla.query.get_or_404(falla_id)

    try:
        template_path = os.path.join(os.path.dirname(__file__), '..', 'templates', 'FORMATO COMITE.docx')
        doc = Document(template_path)

        replacements = {
            '[FECHA]': formatear_fecha_larga(falla.fecha_falta) if falla.fecha_falta else '',
            '[NOMBRE INSTRUCTOR]': falla.nombre_instructor,
            '[NUMERO DECEDULA]': falla.cedula_instructor,
            '[NOMBRE FICHA]': falla.nombre_ficha or '',
            '[NUMERO FICHA]': falla.numero_ficha or '',
            '[NOMBRE APRENDIZ]': falla.nombre_aprendiz,
            '[CEDULA APREDIZ]': falla.documento_aprendiz,
            '[CORREO APRENDIZ]': falla.correo_aprendiz,
            '[TELEFONO APRENDIZ]': falla.telefono_aprendiz,
            'Descripción de Faltas': 'Descripción de Faltas: ' + falla.descripcion_faltas,
            'Evidencia Fotográfica': 'Evidencia Fotográfica',
            'Quien presenta la queja y/o informe por la firma': 'Quien presenta la queja y/o informe por la firma',
        }
        _replace_placeholders_doc(doc, replacements)

        # Agregar evidencias fotográficas
        if falla.evidencia:
            evidencia_paths = falla.evidencia.split(',')
            for paragraph in doc.paragraphs:
                if 'Evidencia Fotográfica' in paragraph.text:
                    paragraph.text = paragraph.text.replace('Evidencia Fotográfica', '')
                    for ev_path in evidencia_paths:
                        if ev_path.strip():
                            full_path = os.path.join(os.path.dirname(__file__), '..', 'static', ev_path.strip())
                            if os.path.exists(full_path):
                                try:
                                    # Agregar cada imagen en un nuevo run para que se muestre
                                    run = paragraph.add_run()
                                    run.add_picture(full_path, width=Inches(4.0))
                                    # Agregar salto de línea y espacio después de cada imagen
                                    paragraph.add_run('\n\n')
                                except Exception as e:
                                    paragraph.add_run(f"[Evidencia: {os.path.basename(full_path)}]\n\n")
                    break

        # Agregar firma
        if falla.firma:
            full_path = os.path.join(os.path.dirname(__file__), '..', 'static', falla.firma.strip())
            if os.path.exists(full_path):
                for paragraph in doc.paragraphs:
                    if 'Quien presenta la queja y/o informe por la firma' in paragraph.text:
                        paragraph.text = paragraph.text.replace('Quien presenta la queja y/o informe por la firma', '')
                        run = paragraph.add_run()
                        try:
                            run.add_picture(full_path, width=Inches(2.0))
                        except Exception:
                            paragraph.text = f"[Firma: {os.path.basename(full_path)}]"
                        break

        import tempfile
        import uuid
        temp_dir = tempfile.gettempdir()
        temp_docx = os.path.join(temp_dir, f'formato_comite_{falla_id}_{uuid.uuid4().hex}.docx')
        temp_pdf = os.path.join(temp_dir, f'formato_comite_{falla_id}_{uuid.uuid4().hex}.pdf')
        doc.save(temp_docx)

        return _enviar_pdf_siempre(
            temp_docx=temp_docx,
            temp_pdf=temp_pdf,
            download_name=f'formato_comite_{falla.numero_ficha}_{falla.nombre_aprendiz.replace(" ", "_")}.pdf',
            doc_para_respaldo=doc,
            titulo_respaldo='Formato comite generado',
        )
    except Exception as e:
        flash(f'Error al generar el formato: {str(e)}', 'error')
        return redirect(url_for('auth.vista_historial'))